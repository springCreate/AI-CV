"""
简历解析服务

支持 PDF / Word 文件解析，结合 AI 完成结构化信息提取。
"""
import os
import logging
from pathlib import Path
from typing import Dict, Any, Optional

from app.services.ai_service import DeepSeekService

logger = logging.getLogger(__name__)


class ResumeParser:
    """简历解析器"""

    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析 PDF 文件为纯文本"""
        import pdfplumber
        text_parts = []
        with pdfplumber.open(file_path) as pdf:
            for i, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                text_parts.append(f"=== 第 {i} 页 ===\n{page_text}")
        return "\n\n".join(text_parts)

    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析 Word 文件为纯文本"""
        from docx import Document
        doc = Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)
        # 表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        return "\n".join(text_parts)

    @staticmethod
    def parse_file(file_path: str, file_type: str) -> str:
        """根据文件类型解析"""
        file_path = str(file_path)
        if file_type.lower() == "pdf":
            return ResumeParser.parse_pdf(file_path)
        elif file_type.lower() in ("docx", "doc"):
            return ResumeParser.parse_docx(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {file_type}")

    @staticmethod
    def extract_structured(text: str) -> Dict[str, Any]:
        """
        调用 AI 将简历纯文本结构化为 JSON。

        返回结构：
        {
            "basic": {"real_name", "gender", "age", "phone", "email", "location"},
            "target": {"target_position", "target_city", "expected_salary"},
            "self_evaluation": str,
            "educations": [{"school", "major", "degree", "start_date", "end_date", "description"}],
            "experiences": [{"company", "position", "job_type", "start_date", "end_date", "description"}],
            "projects": [{"name", "role", "start_date", "end_date", "description", "tech_stack"}],
            "skills": [{"name", "level", "category"}]
        }
        """
        system_prompt = """你是一个专业的简历解析助手。请将用户提供的简历纯文本解析为结构化 JSON 数据。

要求：
1. 严格输出 JSON 格式，不要任何额外说明文字
2. 字段缺失时使用 null 或空字符串
3. 日期统一为 "YYYY-MM" 格式，"至今" 保留为 "至今"
4. 技能 level 取值：了解/熟悉/熟练/精通
5. 技能 category 取值：编程语言/框架/工具/数据库/软技能/其他
6. 经历描述保留原文，不做美化

输出 JSON 结构：
{
  "basic": {"real_name": "", "gender": "", "age": null, "phone": "", "email": "", "location": ""},
  "target": {"target_position": "", "target_city": "", "expected_salary": ""},
  "self_evaluation": "",
  "educations": [{"school": "", "major": "", "degree": "", "start_date": "", "end_date": "", "description": ""}],
  "experiences": [{"company": "", "position": "", "job_type": "", "start_date": "", "end_date": "", "description": ""}],
  "projects": [{"name": "", "role": "", "start_date": "", "end_date": "", "description": "", "tech_stack": ""}],
  "skills": [{"name": "", "level": "", "category": ""}]
}"""

        result = DeepSeekService.chat_json(system_prompt, text, temperature=0.1)
        # 校验必要字段
        result.setdefault("basic", {})
        result.setdefault("target", {})
        result.setdefault("self_evaluation", "")
        result.setdefault("educations", [])
        result.setdefault("experiences", [])
        result.setdefault("projects", [])
        result.setdefault("skills", [])
        return result
